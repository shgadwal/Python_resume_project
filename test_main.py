import pytest
from io import BytesIO
from docx import Document
from PyPDF2 import PdfFileWriter, PdfFileReader
from datetime import datetime
import os

# Import the functions to be tested
from app3 import convert_docx_to_text, convert_pdf_to_text, save_text_to_file, generate_save_path, validate_file_size


# Test convert_docx_to_text function
def test_convert_docx_to_text():
    doc = Document()
    doc.add_paragraph("Hello, World!")
    doc.add_paragraph("This is a test document.")
    # Save the document as bytes
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)  # Reset the stream position

    # Convert the bytes stream to text
    result = convert_docx_to_text(doc_bytes)

    assert "Hello, World!" in result
    assert "This is a test document." in result


# Test convert_pdf_to_text function
def test_convert_pdf_to_text():
    pdf_writer = PdfFileWriter()
    pdf_writer.addPage("Hello, World!")
    pdf_writer.addPage("This is a test PDF.")
    # Save the PDF as bytes
    pdf_bytes = BytesIO()
    pdf_writer.write(pdf_bytes)
    pdf_bytes.seek(0)  # Reset the stream position

    # Convert the bytes stream to text
    result = convert_pdf_to_text(pdf_bytes)

    assert "Hello, World!" in result
    assert "This is a test PDF." in result


# Test save_text_to_file function
def test_save_text_to_file(tmpdir):
    text = "This is a test file."
    filename = os.path.join(tmpdir, "test_file.txt")

    # Save the text to a file
    save_text_to_file(text, filename)

    # Read the saved file
    with open(filename, "r") as f:
        result = f.read()

    assert result == text


# Test generate_save_path function
def test_generate_save_path():
    first_name = "John"
    last_name = "Doe"
    expected_timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    expected_filename = f"{first_name}_{last_name}_{expected_timestamp}.txt"

    # Generate the save path
    save_path = generate_save_path(first_name, last_name)

    assert os.path.basename(save_path) == expected_filename


# Test validate_file_size function
def test_validate_file_size():
    small_file_size = 1000
    large_file_size = 2000000

    # Validate small file size
    assert validate_file_size(small_file_size) is True

    # Validate large file size
    assert validate_file_size(large_file_size) is False


# Run the tests
if __name__ == "__main__":
    pytest.main()
